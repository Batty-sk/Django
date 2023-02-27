import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
import time
from docx2pdf import convert
def Resume_Generator(resume,template_no,has_pic):
    k=0;
    #doc=f"Z://Resume_builder//templates//resume_templates//template{template_no}.docx";
    document=docx.Document(f"L://Django//Django Projects//Resume_builder//templates//resume_templates//template{template_no}.docx")

    for para in document.paragraphs:
        run=para.runs;
        i=0;
        for obj in run:
            print(obj.text)
            obj.text=obj.text.replace(' ','')
            if obj.text in resume.keys():
                print('yes')
                if (type(resume[obj.text])==list or resume[obj.text]==None):
                    print("in List if")
                    #insert before 
                    if (obj.text=='ts' or obj.text=='lls'):
                        print('when obj.txt=',obj.text)
                        if resume[obj.text][k]!='None':
                            obj.text=resume[obj.text][k];
                        else:
                            obj.text='';
                        k+=1;
                        
                    else:
                        for j in resume[obj.text]:
                            if j!='None':
                                para.insert_paragraph_before(j)
                        obj.text=''

                        try:
                            p=para._element
                            p.getparent().remove(p)
                            p._p = p._element = None
                        except:
                            pass
                else:
                    para.runs[i].text=resume[obj.text]
            i+=1
            if k==4:
                k=0;
    document.save("L://Django//Django Projects//Resume_builder//static//template_output//output.docx")
    if has_pic:
        print("has pic")
        add_images("L://Django//Django Projects//Resume_builder//static//template_output//output.docx");
    convert("L://Django//Django Projects//Resume_builder//static//template_output//output.docx","L://Django//Django Projects//Resume_builde2//Resume_builder//static//template_output//output.pdf");

def add_images(doc):
    document=docx.Document(doc)
    for para in document.paragraphs:
        run=para.runs;
        i=0;
        for obj in run:
            if obj.text=='EMAIL:':
                obj.text='';
                obj.add_picture('L://Django//Django Projects//Resume_builder//email.png',width=Inches(0.3))
            elif obj.text=="PHONE:":
                obj.text='';
                obj.add_picture('L://Django//Django Projects//Resume_builder//phone.png',width=Inches(0.3));
        document.save("L://Django//Django Projects//Resume_builder//static//template_output//output.docx")
